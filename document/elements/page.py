from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

from document.elements.misc import add_blue_bar, white_space
from document.elements.border import add_left_border , add_bottom_border, add_top_border, add_field_element

def create_cover_page(doc, properties: dict, withMarkScheme: bool = True):
   # Cover page elements
  add_blue_bar(doc.add_paragraph(), 36)

  white_space(doc, space_before=168)

  # Image placeholder
  img_placeholder = doc.add_paragraph()
  img_run = img_placeholder.add_run()
  img_run.font.size = Pt(10)
  img_run.add_picture("assets/brand_logo.png", width=Inches(1.46))
  img_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER

  # Company Brand
  title = doc.add_paragraph()
  title_run = title.add_run(properties['company_name'])
  title_run.font.name = 'Calibri'
  title_run.font.bold = True
  title_run.font.size = Pt(28)
  title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  title.paragraph_format.space_after = Pt(10)
  
  white_space(doc, space_before=10)
  white_space_with_border = doc.add_paragraph()
  add_bottom_border(white_space_with_border)
  white_space(doc, space_before=20)
  
  subject = doc.add_paragraph()
  subject_run = subject.add_run(properties['subject'])
  subject_run.font.name = 'Calibri'
  subject_run.font.size = Pt(20)
  subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
  subject.paragraph_format.space_after = Pt(10)
  
  topic = doc.add_paragraph()
  topic_run = topic.add_run(properties['topic'])
  topic_run.font.name = 'Calibri'
  topic_run.font.bold = True
  topic_run.font.size = Pt(26)
  topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
  topic.paragraph_format.space_after = Pt(10)
  
  type = doc.add_paragraph()
  type_run = type.add_run(
    f"{f"{properties['paper_and_type']}      " if properties['paper_and_type'] else ""}{"Questions & Mark Scheme" if withMarkScheme else "Questions Only"}"
  )
  type_run.font.name = 'Calibri'
  type_run.font.size = Pt(17)
  type.alignment = WD_ALIGN_PARAGRAPH.CENTER
  type.paragraph_format.space_after = Pt(0)

  for _ in range(10):
    white_space(doc, font_name="Times New Roman", font_size=10)

  add_blue_bar(doc.add_paragraph(), 24)
  
def define_header_and_footer(doc, properties: dict):
  # Headers
  global_header = doc.sections[0].header.paragraphs[0]
  global_header.style = 'Normal'

  header_tab_stops = global_header.paragraph_format.tab_stops
  header_tab_stops.clear_all()
  header_tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT)

  header_logo_run = global_header.add_run()
  header_logo_run.add_picture("assets/brand_logo.png", width=Inches(0.21))

  left_header_run = global_header.add_run(f"  {properties["header_brand"]}")
  left_header_run.font.name = "Calibri"
  left_header_run.font.bold = True
  left_header_run.font.size = Pt(9)
  left_header_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

  right_header_run = global_header.add_run(f"\t{properties['subject']}  ·  {properties['topic']}{f"  ·  {properties['paper_and_type']}" if properties['paper_and_type'] else ""}")
  right_header_run.font.name = "Calibri"
  right_header_run.font.size = Pt(9)
  right_header_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
  
  add_bottom_border(global_header, sz="6")
  
  # Footer
  global_footer = doc.sections[0].footer.paragraphs[0]
  global_footer.paragraph_format.space_after = Pt(0)
  global_footer.style = 'Normal'
  
  footer_tab_stops = global_footer.paragraph_format.tab_stops
  footer_tab_stops.clear_all()
  footer_tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT)
  
  global_footer.add_run(f"© {properties["company_name"]}\t")
  global_footer.add_run("Page ")

  run_page = global_footer.add_run()
  add_field_element(run_page, "PAGE")
  
  global_footer.add_run(" of ")

  run_numpages = global_footer.add_run()
  add_field_element(run_numpages, "NUMPAGES")
  
  # Style all runs in footer
  for run in global_footer.runs:
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
  
  add_top_border(global_footer, sz="6", space="4")


def create_question(doc, path, index):
  # Pages
  question_label = doc.add_paragraph()
  add_left_border(question_label)
  question_label.paragraph_format.page_break_before = True
  question_label.paragraph_format.space_before = Pt(0)
  question_label.paragraph_format.space_after = Pt(0)
  question_label.paragraph_format.left_indent = Pt(10.8)
  question_label_run = question_label.add_run(f"Question {index}")
  question_label_run.font.name = "Calibri"
  question_label_run.font.bold = True
  question_label_run.font.size = Pt(13)
  question_label_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

  # # Put logic of inserting image here
  question_img = doc.add_paragraph()
  question_img.paragraph_format.space_before = Pt(6)
  question_img_run = question_img.add_run()
  question_img_run.font.size = Pt(10)
  question_img_run.add_picture(str(f"{path.parent}\\q{index:02d}_question.png"), width=Inches(6.3))
  
  white_space(doc, space_before=12, isCentered=False)
  
  
def create_mark_scheme(doc, path, index):
  mark_scheme_label = doc.add_paragraph()
  mark_scheme_label.paragraph_format.space_before = Pt(11)
  mark_scheme_label.paragraph_format.space_after = Pt(4)
  mark_scheme_label.paragraph_format.left_indent = Pt(7.2)
  mark_scheme_label_run = mark_scheme_label.add_run(f"Mark Scheme — Question {index}")
  mark_scheme_label_run.font.name = "Calibri"
  mark_scheme_label_run.font.bold = True
  mark_scheme_label_run.font.size = Pt(11)
  mark_scheme_label_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
 
  # For loop here when logic for splitting mark scheme is done
  # Put logic of inserting image here
  mark_scheme_img = doc.add_paragraph()
  mark_scheme_img_run = mark_scheme_img.add_run()
  mark_scheme_img_run.font.size = Pt(10)
  mark_scheme_img_run.add_picture(str(f"{path.parent}\\q{index:02d}_answer.png"), width=Inches(5))
 
  white_space(doc, space_before=12, isCentered=False)